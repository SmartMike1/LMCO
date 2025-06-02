import textwrap
import tkinter.font as tkFont
import webbrowser
from datetime import datetime
from tkinter import *
from tkinter import messagebox, ttk

from docx import Document
from PIL import Image as PILImage
from PIL import ImageTk as PILImageTk
from sqlalchemy import text

from Functions import *


# Каркас инициализации окна
class BasedWindow(Toplevel):
    # x, y - некоторый сдвиг окна, чтобы окна не перекрывались
    # width=700
    def __init__(self, master, engine, title,
                width=750, height=500, x=0, y=0):
        super().__init__(master)
        self.master = master
        self.engine = engine
        self.title(title)
        self.font = ("Arial", 20)
        self.configure(bg="white")
        center_window(self, width, height, x, y)
        self.resizable(False, False)
        # self.resizable(True, True)

# Окно авторизации
class LoginWindow(BasedWindow):
    def __init__(self, master, engine):
        super().__init__(master, engine, title="Авторизация",
                        width=450, height=200)
        
        self.protocol("WM_DELETE_WINDOW", master.destroy)
        self.init_widgets()

    def init_widgets(self):
        LoginCanvas = Canvas(self, bg="white")
        LoginCanvas.pack(fill="both", expand=True)
        LoginCanvas.update() # Обновляем информацию о размерах окна

        label1 = styled_label(self, text="Логин:", font=self.font)
        self.login = Entry(self, font=self.font, background="lightgray")

        label2 = styled_label(self, text="Пароль:", font=self.font)
        self.password = Entry(self, show="*", font=self.font,
                            background="lightgray")

        button1 = styled_button(self, text="Войти", width=8, font=self.font,
                                command=self.try_login)
        button2 = styled_button(self, text="Обновить", width=8, font=self.font,
                                command=update_version)

        LoginCanvas.create_window(65, 50, window=label1)
        LoginCanvas.create_window(65, 100, window=label2)
        LoginCanvas.create_window(275, 50, window=self.login)
        LoginCanvas.create_window(275, 100, window=self.password)
        LoginCanvas.create_window(165, 130, anchor="nw", window=button1)
        # LoginCanvas.create_window(225, 130, anchor="nw", window=button2)

        self.login.focus_set()
        self.login.bind("<FocusIn>", on_focus_in)
        self.password.bind("<FocusIn>", on_focus_in)
        button1.bind("<Button-1>", on_key_press)
        button1.bind("<Button-3>", on_key_press)
        button1.bind("<KeyPress-Return>", on_key_press)
        button1.bind("<KeyRelease-Return>", on_key_release)

        # Изменяем размер виджетов при изменении геометрической сетки
        for i in range(0,2):
            self.columnconfigure(i, weight=1)
            self.rowconfigure(i, weight=1)

    def try_login(self):
        login = self.login.get()
        password = self.password.get()

        try:
            with engine.connect() as conn:
                result = conn.execute(text(
                    "SELECT COUNT(*) FROM Пользователи WHERE Имя=:u AND Пароль=:p"
                ), {'u': login, 'p': password})

                if result.scalar() == 1:
                    self.destroy()
                    app = Application(master=self.master, engine=engine, login=login)
                else:
                    messagebox.showerror("Ошибка", "Неверный логин или пароль")
        except Exception as e:
            print(e)
            messagebox.showerror("Ошибка подключения", str(e))
            self.master.destroy()

# Окно приложения
class Application(BasedWindow):
    # Инициализация окна
    def __init__(self, master=None, engine=None, login=None):
        super().__init__(master, engine, title="Аргентум")
        self.protocol("WM_DELETE_WINDOW", master.destroy)

        # print("Image =", Image)
        # print("type(Image) =", type(Image))
        self.img = PILImage.open("..\images\Logo-LMCO.png")
        self.img = self.img.resize((300, 300))
        self.photo = PILImageTk.PhotoImage(self.img)

        self.resizable(False, False)
        self.font = ("Arial", 20)
        self.login = login

        self.create_widgets()
        set_white_bg_recursive(self)

    # Начальное окно с кнопками действия
    def create_widgets(self):
        self.focus_force()
        self.MainCanvas = Canvas(self, bg="white")
        self.MainCanvas.pack(fill="both", expand=True)

        self.MainCanvas.update() # Обновляем информацию о размерах окна
        self.MainCanvas.create_image(420, 140,
                                image=self.photo, anchor="nw")
        
        text1 = "         Добро пожаловать!\n"
        text2 = "Выберите желаемое действие\n"
        itext = text1 + text2
        text3 = "Настройки Реактивов"

        # Кнопки действия для Реактивов
        self.button_1 = styled_button(self, text="Найти", font=self.font,
                                    width=23, command=self.search_for_R)
        self.button_2 = styled_button(self, text="Добавить", font=self.font,
                                    width=10, command=self.add_new)
        self.button_3 = styled_button(self, text="Изменить", font=self.font,
                                    width=10, command=self.correct_Rinfo)
        self.button_4 = styled_button(self, text="Показать все Реактивы",
                                    font=self.font, width=23,
                                    command=self.read_all_R)
        self.button_5 = styled_button(self, text="Список покупок", font=self.font,
                                    width=23, bg="#69b900", command=self.order_list)
        self.button_6 = styled_button(self, text="Другие функции",
                                    font=self.font, width=23,
                                    command=self.user_settings)
        
        self.MainCanvas.create_text(200, 25, text=itext, font=self.font,
                                anchor="nw")
        self.MainCanvas.create_text(25, 100, text=text3, font=self.font,
                                anchor="nw")
        self.MainCanvas.create_window(25, 140, anchor="nw", window=self.button_1)
        self.MainCanvas.create_window(25, 200, anchor="nw", window=self.button_2)
        self.MainCanvas.create_window(233, 200, anchor="nw", window=self.button_3)
        self.MainCanvas.create_window(25, 260, anchor="nw", window=self.button_4)
        self.MainCanvas.create_window(25, 340, anchor="nw", window=self.button_5)
        self.MainCanvas.create_window(25, 400, anchor="nw", window=self.button_6)
    
    # Окно настроек Пользователей
    def user_settings(self):
        window = BasedWindow(master=self.master, engine=engine,
                        title="Дополнительные настройки", x=50, y=50)
        window.focus_force()
        wCanvas = Canvas(window, bg="white")
        wCanvas.pack(fill="both", expand=True)

        wCanvas.create_image(420, 140, image=self.photo, anchor="nw")

        text1 = "Выберите желаемое действие"
        text2 = "Настройки Пользователей"
        text3 = "Другие действия"

        # Кнопки действия для Пользователей
        self.button_7 = styled_button(wCanvas, text="Найти", font=self.font,
                                    width=23, command=self.search_for_U)
        self.button_8 = styled_button(wCanvas, text="Добавить", font=self.font,
                                    width=10, command=self.create_user)
        self.button_9 = styled_button(wCanvas, text="Изменить", font=self.font,
                                    width=10, command=self.correct_Uinfo)
        self.button_10 = styled_button(wCanvas, text="Показать всех",
                                    font=self.font, width=23,
                                    command=self.read_all_users)
        
        # Другие действия
        self.button_11 = styled_button(wCanvas, text="Журнал изменений",
                                font=self.font, width=23, bg="#69b900",
                                command=self.read_log)

        wCanvas.create_text(200, 25, text=text1, font=self.font, anchor="nw")
        wCanvas.create_text(25, 65, text=text2, font=self.font, anchor="nw")
        wCanvas.create_text(25, 300, text=text3, font=self.font, anchor="nw")
        wCanvas.create_window(25, 100, anchor="nw", window=self.button_7)
        wCanvas.create_window(25, 160, anchor="nw", window=self.button_8)
        wCanvas.create_window(233, 160, anchor="nw", window=self.button_9)
        wCanvas.create_window(25, 220, anchor="nw", window=self.button_10)
        wCanvas.create_window(25, 335, anchor="nw", window=self.button_11)

    # Универсальная функция для отображения всех реактивов/пользователей
    def read_all(self, proc_name, window_title):
        window = BasedWindow(master=self.master, engine=engine,
                            title=window_title,
                            width=700, x=50, y=50)
        window.focus_force()

        frame = Frame(window, bg="white")
        frame.pack(fill="both", expand=True)

        get_table(frame, proc_name, frame.winfo_width())

    # Универсальная функция для окна поиска всех реактивов/пользователей
    def search_for(self, title, label_text, func):
        self.window = BasedWindow(master=self.master, engine=engine,
                            title=title, width=800, x=50, y=50)
    
        self.window.focus_force()
        self.window.update()
    
        self.form = Frame(self.window, width=self.window.winfo_width(), bg="white")
        self.form.grid(row=0, column=0, columnspan=3, sticky="snwe", padx=10, pady=10)

        label = styled_label(self.form, text=label_text, font=self.font)
        label.grid(row=0, column=0, sticky="w", padx=10, pady=10)

        self.number = Entry(self.form, font=self.font, background="lightgray")
        self.number.grid(row=0, column=1, sticky="w", padx=10, pady=10)

        button = styled_button(self.form, text="Найти", font=self.font, command=func)
        button.grid(row=0, column=2, sticky="w", padx=10, pady=10)

        self.form2 = Frame(self.window, width=self.window.winfo_width(),
                            height=350, bg="white")
        self.form2.grid(row=1, column=0, columnspan=3, sticky="snwe", padx=10, pady=10)

        self.number.bind("<FocusIn>", on_focus_in)

    # Универсальная функция для поиска всех реактивов/пользователей
    def search(self, value, container, proc):
        # Очистка предыдущих виджетов
        for widget in container.winfo_children():
            widget.destroy()

        # Проверка на пустой ввод
        querry = f"{proc} '{value}'"
        result = empty_data(value, text(querry))

        if result is None:
            text1 = "Ничего не найдено\n"
            text2 = "Проверьте введённые данные"
            label = styled_label(container, text=text1 + text2, font=self.font)
            label.grid()
        else:
            get_table(container, querry, self.window.winfo_width() - 25)

    # Функции работы с Реактивами
    def read_all_R(self):
        self.read_all("ReadAllReagents", "Просмотр всех Реактивов")

    def add_new(self):
        window = BasedWindow(master=self.master, engine=engine,
                        title="Добавление нового Реактива", x=50, y=50)
        window.focus_force()

        self.TableName = "Реагенты"
        self.OperationName = "INSERT"

        self.canvas = Canvas(window, width=730, height=400, bg="white")
        self.canvas.grid(row=0, column=0)
        canvas_frame = Frame(self.canvas, bg="white")
        self.canvas.create_window((0, 0), window=canvas_frame, anchor="nw")

        # Создаем Scrollbar и привязываем его к Canvas
        scrollbar = Scrollbar(window, bg="white", orient="vertical",
                                    command=self.canvas.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.canvas.config(yscrollcommand=scrollbar.set)

        # Привязываем событие колесика мыши к прокрутке Canvas
        self.canvas.bind_all("<MouseWheel>", self.on_mouse_wheel)

        try:
            querry = f"ReadAllReagents"
            columns, rows = get_data_for_text(querry)
            old = "КороткоеНазваниеЕдИзм"
            new = "ЕдиницаИзмерения"

            if old in columns:
                index = columns.index(old)
                columns[index] = new

            self.text_fields = []
            self.original_values = []
            self.column_names = []

            for index, col_name in enumerate(columns):
                label = styled_label(canvas_frame, text=col_name, font=self.font)
                label.grid(row=index, column=0, padx=5, pady=5)

                if col_name == "ЕдиницаИзмерения":
                    querryM = f"AllMeasures"
                    columnsM, rowsM = get_data_for_text(querryM)

                    box_widget = ttk.Combobox(canvas_frame, values=rowsM,
                                            state="readonly", font=self.font,
                                            width=34)
                    box_widget.grid(row=index, column=1, padx=5, pady=5)
                    box_widget.set("Выберите из...")
                    self.text_fields.append(box_widget)
                elif col_name == "СтатусРеагента":
                    querryS = f"AllStatus"
                    columns, rowsS = get_data_for_text(querryS)

                    box_widget = ttk.Combobox(canvas_frame, values=rowsS,
                                            state="readonly", font=self.font,
                                            width=34)
                    box_widget.grid(row=index, column=1, padx=5, pady=5)
                    box_widget.set("Выберите из...")
                    self.text_fields.append(box_widget)
                else:
                    text_widget = Text(canvas_frame, width=35, height=2,
                                    font=self.font, relief="solid", wrap=WORD)
                    text_widget.grid(row=index, column=1, padx=5, pady=5)
                    self.text_fields.append(text_widget)

                self.original_values.append("")
                self.column_names.append(col_name)

            # Обновляем, чтобы прокрутка работала правильно
            canvas_frame.update_idletasks()
            self.canvas.config(scrollregion=self.canvas.bbox("all"))

            button = styled_button(window, text="Сохранить Информацию",
                            font=self.font, command=self.insert_newR)
            button.grid(row=2, column=0, sticky="we", padx=45, pady=25)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось выполнить процедуру:\n{e}")
        set_white_bg_recursive(self)

    def insert_newR(self):
        changes = {}

        for i, widget in enumerate(self.text_fields):
            if isinstance(widget, ttk.Combobox):
                new_value = widget.get()
            else:
                new_value = widget.get("1.0", END).strip()
            old_value = self.original_values[i].strip()

            changes[self.column_names[i]] = (old_value, new_value)

        if not changes:
            messagebox.showinfo("Информация", "Нет информации для вставки.")
            return
        
        try:
            querry = text(" SELECT COUNT(*) FROM Реагенты")
            insertR = """INSERT INTO Реагенты (НазваниеРеагента,
                        CAS, МестоНаСкладе, ВнешнийВид, КлассСоединения,
                        Примечание, Ссылка, Формула)
                    VALUES (:НазваниеРеагента,
                        :CAS, :МестоНаСкладе, :ВнешнийВид, :КлассСоединения,
                        :Примечание, :Ссылка, :Формула)"""
            insertC = """INSERT INTO СловарьКоличества (КодРеагента,
                        Количество, КодЕдИзмерения, КодСтатусаРеагента)
                    VALUES (:КодРеагента, :Количество,
                        :КороткоеНазваниеЕдИзм, :СтатусРеагента)"""
            
            keys_reagents = ['НазваниеРеагента', 'CAS', 'МестоНаСкладе',
                'ВнешнийВид', 'КлассСоединения', 'Примечание', 'Ссылка',
                'Формула']
            keys_count = ['КодРеагента', 'Количество',
                        'КороткоеНазваниеЕдИзм', 'СтатусРеагента']
            
            reagents_dict = {k: changes[k][1] for k in keys_reagents
                            if k in changes}
            quantity_dict = {k: changes[k][1] for k in changes
                            if k in keys_count}
            quantity_dict['Количество'] = int(quantity_dict['Количество'])

            result = replace_values_with_keys(quantity_dict)

            with engine.connect() as conn:
                conn.execute(text(insertR), reagents_dict)
                conn.commit()
            
            with engine.connect() as conn:
                cursor = conn.execute(querry)
                count = cursor.fetchone()[0]
                result['КодРеагента'] = count+1

                conn.execute(text(insertC), result)
                conn.commit()

            messagebox.showinfo("Успех", "Новый Реактив добавлен!")

        except Exception as e:
            print(e)
            messagebox.showerror("Ошибка", f"Не удалось выполнить процедуру:\n{e}")

    def search_for_R(self):
        self.search_for(title="Поиск Реактива", label_text="CAS / Название:",
                        func=self.search_R)

    def search_R(self):
        number = self.number.get()
        self.search(number, self.form2, "SearchReagent")
        
    def correct_Rinfo(self):
        self.Cwindow = BasedWindow(master=self.master, engine=engine,
                title="Изменение данных по Реактиву",
                width=800, x=50, y=50)

        self.Cwindow.update()
        self.Cwindow.focus_force()

        self.TableName = "Реагенты"
        self.OperationName = "UPDATE"

        # Создаем Canvas
        self.canvas = Canvas(self.Cwindow, width=780, height=325, bg="white")
        self.canvas.grid(row=1, column=0)

        # Создаем Scrollbar и привязываем его к Canvas
        self.scrollbar = Scrollbar(self.Cwindow, orient="vertical",
                                    command=self.canvas.yview, bg="white")
        self.scrollbar.grid(row=1, column=1, sticky="ns")
        self.canvas.config(yscrollcommand=self.scrollbar.set)

        # Привязываем событие колесика мыши к прокрутке Canvas
        self.canvas.bind_all("<MouseWheel>", self.on_mouse_wheel)

        # Создаем Frame внутри Canvas для размещения содержимого
        self.canvas_frame = Frame(self.canvas, bg="white")
        self.canvas.create_window((0, 0), window=self.canvas_frame, anchor="nw")

        self.form = Frame(self.Cwindow, bg="white")
        self.form.grid(row=0, column=0, columnspan=3, sticky="snwe",
                        padx=10, pady=10)

        label = styled_label(self.form, text="CAS / Название", font=self.font)
        label.grid(row=0, column=0, sticky="w", padx=10, pady=10)
        self.number = Entry(self.form, font=self.font, background="lightgray")
        self.number.grid(row=0, column=1, sticky="w", padx=10, pady=10)

        button = styled_button(self.form, text="Найти", font=self.font, command=self.choose)
        button.grid(row=0, column=2, sticky="w", padx=10, pady=10)

        self.number.bind("<FocusIn>", on_focus_in)
        button.bind("<Return>")

    def on_mouse_wheel(self, event):
        self.canvas.yview_scroll(-1 * (event.delta // 120), "units")
        # self.canvas1.yview_scroll(-1 * (event.delta // 120), "units")

    def choose(self):
        self.window = BasedWindow(master=self.master, engine=engine,
                title="Выбор нужного Реактива",
                width=620, x=50, y=50)

        number = self.number.get()

        # Проверка на пустой запрос
        querry = f"SearchReagent '{number}'"
        result = empty_data(number, text(querry))

        # Отображаем данные
        if number == "" or result == None:
            self.window.destroy()
            text1 = "Ничего не найдено\n"
            text2 = "Проверьте введённые данные"
            self.canvas.create_text(200, 50, text=text1+text2, font=self.font)
        else:
            self.sheet = get_table(self.window, querry,
                                self.window.winfo_width() - 25)
            self.sheet.enable_bindings(("cell_select", "double_click"))
            self.sheet.bind("<Double-Button-1>", self.on_sheet_double_click)

    def on_sheet_double_click(self, event):
        self.Cwindow.focus_force()
        # Удаляем старые виджеты из контейнера (если есть)
        for widget in self.canvas_frame.winfo_children():
            widget.destroy()

        row_index = self.sheet.get_currently_selected()[0]
        if row_index is None:
            return

        row_values = self.sheet.get_row_data(row_index)
        self.reagent_id = row_values[0]

        try:
            querry = f"CorrectReagentInformation '{self.reagent_id}'"
            columns, rows = get_data_for_text(querry)

            old = "КороткоеНазваниеЕдИзм"
            new = "ЕдиницаИзмерения"

            if old in columns:
                index = columns.index(old)
                columns[index] = new

            self.text_fields = []
            self.original_values = []
            self.column_names = []

            # print(len(rows), len(columns))
            for i in range (len(rows)):
                for index, col_name in enumerate(columns):
                    label = styled_label(self.canvas_frame, text=col_name, font=self.font)
                    label.grid(row=index, column=0, padx=5, pady=5)

                    if col_name == "ЕдиницаИзмерения":
                        querryM = f"AllMeasures"
                        columnsM, rowsM = get_data_for_text(querryM)

                        box_widget = ttk.Combobox(self.canvas_frame, values=rowsM,
                                            state="readonly", font=self.font,
                                            width=34)
                        box_widget.grid(row=index, column=1, padx=5, pady=5)
                    elif col_name == "СтатусРеагента":
                        querryS = f"AllStatus"
                        columnsM, rowsS = get_data_for_text(querryS)

                        box_widget = ttk.Combobox(self.canvas_frame, values=rowsS,
                                            state="readonly", font=self.font,
                                            width=34)
                        box_widget.grid(row=index, column=1, padx=5, pady=5)
                    else:
                        text_widget = Text(self.canvas_frame, width=35, height=2,
                                    font=self.font, relief="solid", wrap=WORD)
                        text_widget.grid(row=index, column=1, padx=5, pady=5)
                        text_widget.insert(1.0, rows[i][index])

                        self.text_fields.append(text_widget)
                        self.original_values.append(rows[i][index])
                        self.column_names.append(col_name)

            # Обновляем, чтобы прокрутка работала правильно
            self.canvas_frame.update_idletasks()
            self.canvas.config(scrollregion=self.canvas.bbox("all"))

            # Закрываем окно "Выбор нужного Реактива"
            self.window.destroy()

            button = styled_button(self.Cwindow, text="Сохранить Информацию",
                                font=self.font, command=self.update_changes)
            button.grid(row=2, column=0, sticky="we", padx=10, pady=10)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось выполнить процедуру:\n{e}")
    
    def update_changes(self):
        changes = {}

        for i, text_widget in enumerate(self.text_fields):
            new_value = text_widget.get("1.0", END).strip()
            old_value = self.original_values[i].strip()

            if new_value != old_value:
                changes[self.column_names[i]] = (old_value, new_value)

        if not changes:
            messagebox.showinfo("Информация", "Нет изменений для сохранения.")
            return
        
        try:
            update   = text("ProcUpdateField :id, :field, :value")
            user     = text("GetUserIdByLogin :login")
            log_proc = text("LogChange :user_id, :table_name, :operation, :field, :before, :after")

            for column, (old, new) in changes.items():
                with engine.connect() as conn:
                    # Определяем ID по Фамилии
                    result = conn.execute(user, {"login": self.login}).fetchone()
                    user_id = result[0]

                    # Изменяем данные Реагента
                    conn.execute(update, {"id": self.reagent_id, "field": column, "value": new})

                    # Логируем изменения
                    conn.execute(log_proc, {"user_id": user_id, "table_name": self.TableName,
                                            "operation": self.OperationName, "field": column,
                                            "before": old, "after": new})
                    conn.commit()

            messagebox.showinfo("Успех", "Изменения успешно сохранены.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить изменения:\n{e}")

    def order_list(self, output_dir="."):
        try:
            today_str = datetime.now().strftime("%Y-%m-%d")
            filename = f"Список_реагентов_{today_str}.docx"
            output_path = f"{output_dir}/{filename}"

            doc = Document()
            doc.add_heading('Список реагентов для заказа', 0)

            querry = f"OrderList"
            columns, rows = get_data_for_text(querry)
            # print(columns, rows)
            if not rows:
                doc.add_paragraph("Нет данных для отображения.")
            else:
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'

            # Заголовки
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(columns):
                hdr_cells[i].text = str(col_name)

            # Данные
            # print(rows)
            for row in rows:
                # print(row)
                row_cells = table.add_row().cells
                for i in range(len(columns)):
                    value = row[i] if i < len(row) else ""
                    row_cells[i].text = str(value)

            doc.save(output_path)
            messagebox.showinfo("Успех", f"Документ успешно сохранён: {output_path}")
        
        except Exception as e:
            print(e)
            messagebox.showerror("Ошибка", f"Не удалось сгенерировать файл:\n{e}")
    
    # !!!!!!! Не написано !!!!!!!
    def delete_reagent(self):
        window = BasedWindow(master=self.master, engine=engine,
                        title="Списание Реактива",
                        width=620, x=50, y=50)
    # !!!!!!! Не написано !!!!!!!

    # Функции работы с Пользователями
    def read_all_users(self):
        self.read_all("ReadAllUsers", "Просмотр всех Пользователей")

    def create_user(self):
        window = BasedWindow(master=self.master, engine=engine,
                        title="Добавление Пользователя", x=50, y=50)
        window.focus_force()

        self.TableName = "Пользователи"

        self.canvas = Canvas(window, width=730, height=400, bg="white")
        self.canvas.grid(row=0, column=0)
        canvas_frame = Frame(self.canvas, bg="white")
        self.canvas.create_window((0, 0), window=canvas_frame, anchor="nw")

        # Создаем Scrollbar и привязываем его к Canvas
        scrollbar = Scrollbar(window, bg="white", orient="vertical",
                                    command=self.canvas.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.canvas.config(yscrollcommand=scrollbar.set)

        # Привязываем событие колесика мыши к прокрутке Canvas
        self.canvas.bind_all("<MouseWheel>", self.on_mouse_wheel)

        try:
            querry = f"ReadAllUsers"
            columns, rows = get_data_for_text(querry)

            self.text_fields = []
            self.original_values = []
            self.column_names = []

            for index, col_name in enumerate(columns):
                label = styled_label(canvas_frame, text=col_name, font=self.font)
                label.grid(row=index, column=0, padx=5, pady=5)

                if col_name == "Статус":
                    status = ["Лаборант", "Управляющий", "Начальник", "Разработчик"]

                    box_widget = ttk.Combobox(canvas_frame, values=status,
                                            state="readonly", font=self.font,
                                            width=34)
                    box_widget.grid(row=index, column=1, padx=5, pady=5)
                    box_widget.set("Выберите из...")
                    self.text_fields.append(box_widget)
                else:
                    text_widget = Text(canvas_frame, width=35, height=2,
                    font=self.font, relief="solid", wrap=WORD)
                    text_widget.grid(row=index, column=1, padx=5, pady=5)
                    # text_widget.insert(1.0, rows[i][index])
                    self.text_fields.append(text_widget)
                    
                self.original_values.append("")
                self.column_names.append(col_name)

            # Обновляем, чтобы прокрутка работала правильно
            canvas_frame.update_idletasks()
            self.canvas.config(scrollregion=self.canvas.bbox("all"))

            button = styled_button(window, text="Сохранить Информацию",
                        font=self.font, command=self.insert_newU)
            button.grid(row=2, column=0, sticky="we", padx=45, pady=25)
        
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось выполнить процедуру:\n{e}")
        set_white_bg_recursive(self)

    def insert_newU(self):
        changes = {}

        for i, widget in enumerate(self.text_fields):
            if isinstance(widget, ttk.Combobox):
                new_value = widget.get()
            else:
                new_value = widget.get("1.0", END).strip()
            changes[self.column_names[i]] = new_value

        if not changes:
            messagebox.showinfo("Информация", "Нет информации для вставки.")
            return
        
        try:
            insertU = """INSERT INTO Пользователи (Имя, Фамилия,
                                                Пароль, Статус)
                    VALUES (:Имя, :Фамилия, :Пароль, :Статус)"""

            with engine.connect() as conn:
                conn.execute(text(insertU), changes)
                conn.commit()

            messagebox.showinfo("Успех", "Новый Пользователь добавлен!")
        except Exception as e:
            print(e)
            messagebox.showerror("Ошибка", f"Не удалось выполнить процедуру:\n{e}")
    
    def search_for_U(self):
        self.search_for(title="Поиск Пользователя", label_text="Имя / Фамилия:",
                        func=self.search_U)

    def search_U(self):
        querry = self.number.get()
        self.search(querry, self.form2, "SearchUser")

    def correct_Uinfo(self):
        self.CUwindow = BasedWindow(master=self.master, engine=engine,
                title="Изменение данных Пользователя",
                width=800, x=50, y=50)

        self.CUwindow.update()
        self.CUwindow.focus_force()

        self.TableName = "Пользователи"
        self.OperationName = "UPDATE"

        self.Ucanvas = Canvas(self.CUwindow, width=780, height=325, bg="white")
        self.Ucanvas.grid(row=1, column=0)

        self.scrollbar = Scrollbar(self.CUwindow, orient="vertical",
                                    command=self.Ucanvas.yview, bg="white")
        self.scrollbar.grid(row=1, column=1, sticky="ns")
        self.Ucanvas.config(yscrollcommand=self.scrollbar.set)

        self.Ucanvas.bind_all("<MouseWheel>", self.on_mouse_wheel_1)

        self.Ucanvas_frame = Frame(self.Ucanvas, bg="white")
        self.Ucanvas.create_window((0, 0), window=self.Ucanvas_frame, anchor="nw")

        self.form = Frame(self.CUwindow, bg="white")
        self.form.grid(row=0, column=0, columnspan=3, sticky="snwe",
                        padx=10, pady=10)

        label = styled_label(self.form, text="Имя / Фамилия", font=self.font)
        label.grid(row=0, column=0, sticky="w", padx=10, pady=10)
        self.number = Entry(self.form, font=self.font, background="lightgray")
        self.number.grid(row=0, column=1, sticky="w", padx=10, pady=10)

        button = styled_button(self.form, text="Найти", font=self.font, command=self.choose_1)
        button.grid(row=0, column=2, sticky="w", padx=10, pady=10)

        self.number.bind("<FocusIn>", on_focus_in)
        button.bind("<Return>")

    def on_mouse_wheel_1(self, event):
        self.Ucanvas.yview_scroll(-1 * (event.delta // 120), "units")
        # self.canvas1.yview_scroll(-1 * (event.delta // 120), "units")

    def choose_1(self):
        self.window = BasedWindow(master=self.master, engine=engine,
                title="Выбор нужного Пользователя",
                width=620, x=50, y=50)

        number = self.number.get()

        # Проверка на пустой запрос
        querry = f"SearchUser '{number}'"
        result = empty_data(number, text(querry))

        # Отображаем данные
        if number == "" or result == None:
            self.window.destroy()
            text1 = "Ничего не найдено\nПроверьте введённые данные"
            self.Ucanvas.create_text(200, 50, text=text1, font=self.font)
        else:
            self.sheet = get_table(self.window, querry,
                                self.window.winfo_width() - 25)
            self.sheet.enable_bindings(("cell_select", "double_click"))
            self.sheet.bind("<Double-Button-1>", self.double_click_1)

    def double_click_1(self, event):
        self.CUwindow.focus_force()
        # Удаляем старые виджеты из контейнера (если есть)
        for widget in self.Ucanvas_frame.winfo_children():
            widget.destroy()

        row_index = self.sheet.get_currently_selected()[0]
        if row_index is None:
            return

        row_values = self.sheet.get_row_data(row_index)
        self.user = row_values[0].strip()
        print(self.user)

        try:
            querry = f"CorrectUserInformation_1 '{self.user}'"
            columns, rows = get_data_for_text(querry)
            self.user_id = rows[0][0]
            columns = columns[1:]
            rows[0] = rows[0][1:]

            print(querry)
            print(columns)
            print(rows)

            self.text_fields = []
            self.original_values = []
            self.column_names = []

            for i in range (len(rows)):
                for index, col_name in enumerate(columns):
                    label = styled_label(self.Ucanvas_frame, text=col_name, font=self.font)
                    label.grid(row=index, column=0, padx=5, pady=5)

                    if col_name == "Статус":
                        status = ["Лаборант", "Управляющий", "Начальник", "Разработчик"]

                        box_widget = ttk.Combobox(self.Ucanvas_frame, values=status,
                                            state="readonly", font=self.font,
                                            width=34)
                        box_widget.grid(row=index, column=1, padx=5, pady=5)
                        box_widget.set("Выберите из...")
                        self.text_fields.append(box_widget)
                    else:
                        text_widget = Text(self.Ucanvas_frame, width=35, height=2,
                                    font=self.font, relief="solid", wrap=WORD)
                        text_widget.grid(row=index, column=1, padx=5, pady=5)
                        text_widget.insert(1.0, rows[i][index])

                        self.text_fields.append(text_widget)

                    self.original_values.append(rows[i][index])
                    self.column_names.append(col_name)

            # Обновляем, чтобы прокрутка работала правильно
            self.Ucanvas_frame.update_idletasks()
            self.Ucanvas.config(scrollregion=self.Ucanvas.bbox("all"))
            self.window.destroy()

            button = styled_button(self.CUwindow, text="Сохранить Информацию",
                                font=self.font, command=self.update_Uchanges)
            button.grid(row=2, column=0, sticky="we", padx=10, pady=10)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось выполнить процедуру:\n{e}")

    def update_Uchanges(self):
        changes = {}

        for i, widget in enumerate(self.text_fields):
            if isinstance(widget, ttk.Combobox):
                new_value = widget.get()
            else:
                new_value = widget.get("1.0", END).strip()
            old_value = self.original_values[i].strip()

            if new_value != old_value:
                changes[self.column_names[i]] = (old_value, new_value)

        if not changes:
            messagebox.showinfo("Информация", "Нет изменений для сохранения.")
            return

        try:
            update   = text("ProcUpdateFieldU :id, :field, :value")
            user     = text("GetUserIdByLogin :login")
            log_proc = text("LogChange :user_id, :table_name, :operation, :field, :before, :after")

            for column, (old, new) in changes.items():
                with engine.connect() as conn:
                    # Определяем ID по Фамилии
                    result = conn.execute(user, {"login": self.login}).fetchone()
                    current_user_id = result[0]

                    # Изменяем данные Пользователя
                    conn.execute(update, {"id": self.user_id, "field": column, "value": new})

                    # Логируем изменения
                    conn.execute(log_proc, {"user_id": current_user_id, "table_name": self.TableName,
                                            "operation": self.OperationName, "field": column,
                                            "before": old, "after": new})
                    conn.commit()

            messagebox.showinfo("Успех", "Изменения успешно сохранены.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить изменения:\n{e}")

    def read_log(self):
        window = BasedWindow(master=self.master, engine=engine,
                            title="Просмотр Журнала Изменений",
                            width=700, x=50, y=50)
        window.focus_force()

        frame = Frame(window, bg="white")
        frame.pack(fill="both", expand=True)

        get_table(frame, "ReadLog", frame.winfo_width())

    # !!!!!!! Не написано !!!!!!!
    def delete_user(self):
        window = BasedWindow(master=self.master, engine=engine,
                        title="Удаление пользователя",
                        width=620, x=50, y=50)
        window.focus_force()

        self.TableName = "Пользователи"
    # !!!!!!! Не написано !!!!!!!

# Запускаемся
if __name__ == "__main__":
    # engine = get_engine()

    root = Tk()
    root.withdraw()
    login = LoginWindow(master=root, engine=engine)

    login.mainloop()